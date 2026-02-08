"""
Word/Excel转换服务
"""
import logging
from pathlib import Path
from typing import List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import openpyxl
from openpyxl.styles import Font, Alignment
import pandas as pd

logger = logging.getLogger(__name__)


class OfficeService:
    """Office文档服务"""

    def __init__(self):
        pass

    def export_to_word(
        self,
        text: str,
        output_path: str = None,
        title: str = "文档"
    ) -> str:
        """
        导出为Word文档

        Args:
            text: 文字内容
            output_path: 输出文件路径
            title: 文档标题

        Returns:
            Word文件路径
        """
        try:
            # 创建Word文档
            doc = Document()

            # 添加标题
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 0, 0)
            title_para.alignment = 1  # 居中

            # 添加空行
            doc.add_paragraph()

            # 添加内容
            content_para = doc.add_paragraph(text)
            content_para.paragraph_format.line_spacing = 1.5
            content_run = content_para.runs[0]
            content_run.font.size = Pt(12)

            # 保存
            if output_path is None:
                output_path = f"document_{timestamp()}.docx"

            doc.save(output_path)

            logger.info(f"Word文档导出成功: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Word文档导出失败: {e}")
            raise Exception(f"Word文档导出失败: {str(e)}")

    def export_to_excel(
        self,
        texts: List[str],
        output_path: str = None,
        sheet_name: str = "Sheet1"
    ) -> str:
        """
        导出为Excel文档

        Args:
            texts: 文字列表
            output_path: 输出文件路径
            sheet_name: 工作表名称

        Returns:
            Excel文件路径
        """
        try:
            # 创建Excel工作簿
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = sheet_name

            # 写入表头
            ws['A1'] = "序号"
            ws['B1'] = "内容"
            ws['C1'] = "时间"

            # 设置表头样式
            header_font = Font(size=12, bold=True, color="FFFFFF")
            header_fill = openpyxl.styles.PatternFill(
                start_color="4472C4",
                end_color="4472C4",
                fill_type="solid"
            )

            for cell in ws[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")

            # 写入数据
            from datetime import datetime

            for i, text in enumerate(texts, start=1):
                ws[f'A{i+1}'] = i
                ws[f'B{i+1}'] = text
                ws[f'C{i+1}'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

                # 设置列宽
                ws.column_dimensions['A'].width = 8
                ws.column_dimensions['B'].width = 50
                ws.column_dimensions['C'].width = 20

                # 设置数据样式
                data_font = Font(size=11)
                for col in ['A', 'B', 'C']:
                    cell = ws[f'{col}{i+1}']
                    cell.font = data_font
                    cell.alignment = Alignment(
                        horizontal="left",
                        vertical="center",
                        wrap_text=True
                    )

            # 保存
            if output_path is None:
                output_path = f"document_{timestamp()}.xlsx"

            wb.save(output_path)

            logger.info(f"Excel文档导出成功: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Excel文档导出失败: {e}")
            raise Exception(f"Excel文档导出失败: {str(e)}")

    def export_table_to_excel(
        self,
        data: List[dict],
        output_path: str = None,
        sheet_name: str = "Sheet1"
    ) -> str:
        """
        导出表格数据到Excel

        Args:
            data: 表格数据（字典列表）
            output_path: 输出文件路径
            sheet_name: 工作表名称

        Returns:
            Excel文件路径
        """
        try:
            # 使用pandas
            df = pd.DataFrame(data)

            # 保存
            if output_path is None:
                output_path = f"table_{timestamp()}.xlsx"

            df.to_excel(output_path, sheet_name=sheet_name, index=False)

            logger.info(f"Excel表格导出成功: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Excel表格导出失败: {e}")
            raise Exception(f"Excel表格导出失败: {str(e)}")


def timestamp():
    """生成时间戳"""
    from datetime import datetime
    return datetime.now().strftime("%Y%m%d_%H%M%S")
